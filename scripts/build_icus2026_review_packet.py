from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


PROJECT_ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = PROJECT_ROOT / "docs" / "submission"
OUT_DOCX = OUT_DIR / "icus2026_ozet_bildiri_degerlendirme_paketi.docx"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(30, 30, 30)
MUTED = RGBColor(90, 90, 90)
LIGHT_FILL = "F2F4F7"
CALLOUT_FILL = "F4F6F9"
BORDER = "D9E2EC"


TURKISH_ABSTRACT = """
İklim değişikliğine dirençli kentler yalnızca afet tehlikelerinin dağılımı üzerinden değil, bu tehlikelerin hangi kentsel dokular içinde karşılandığı üzerinden de değerlendirilmelidir. Sokak ağının geçirgenliği, yapı adalarının büyüklüğü, yapı yoğunluğu, açıklık oranı, gölge ve güneşlenme koşulları, donatı erişimi ve sosyal kırılganlık göstergeleri, farklı mahalle parçalarının iklim baskılarına verdiği yanıtı doğrudan etkiler. Buna karşın birçok kentsel dirençlilik değerlendirmesi idari sınırlar, genel risk endeksleri veya tekil tehlike katmanları düzeyinde kalmakta, kentsel tasarım ölçeğinde okunabilen doku farklarını yeterince görünür kılmamaktadır. Bu eksiklik, özellikle kıyı, yamaç, tarihî merkez, sanayi kenarı ve yeni gelişme dokularının yan yana bulunduğu İzmir gibi karmaşık metropoliten alanlarda daha belirgin hale gelmektedir.

Bu çalışma, İzmir Körfezi çevresinde seçilecek temsilî kentsel doku örnekleri üzerinden sokak bazlı bir morfoloji ve dirençlilik analitiği iş akışı geliştirmeyi amaçlamaktadır. Çalışmanın temel sorusu, farklı sokak-doku tiplerinin erişilebilirlik, mikroiklim maruziyeti, sosyal kırılganlık ve iklim uyum önceliği bakımından nasıl ayrıştığıdır. Bu kapsamda tarihî merkez dokusu, planlı ızgara doku, orta ve yüksek yoğunluklu apartman blokları, kıyı dönüşüm alanları, sanayi-lojistik kenarlar, yamaç yerleşimleri ve yeni gelişme alanları gibi karşılaştırılabilir örneklerin ele alınması öngörülmektedir. Örnek alanlar 400 veya 800 metrelik yürüme yakalamaları, sokak koridoru tamponları ya da eşdeğer alanlı morfoloji ızgaraları üzerinden standartlaştırılacaktır. Böylece farklı büyüklük ve konumdaki parçalar ortak bir ölçüm zemini üzerinde karşılaştırılabilecektir. Örnek seçimi, yalnızca mekânsal dağılımı değil, doku karakterini de temsil edecek biçimde yapılacak; her örnek alanın sokak sürekliliği, yapı adası mantığı, açık alan ilişkisi ve gündelik erişim potansiyeli birlikte ele alınacaktır.

Yöntem, açık kaynaklı QGIS ortamında yürütülecek PlanX Urban Resilience odaklı bir iş akışına dayanmaktadır. Analitik omurgada sokak ağı hazırlama, segment tabanlı mekânsal dizim ölçümleri, ağ merkeziliği, yapı formu metrikleri, morfolojik tessellation, Spacematrix yoğunluk göstergeleri, sokak yönelim düzeni, çoklu donatı erişimi, ısı adası riski, sosyal kırılganlık ve acil durum erişilebilirliği gibi göstergeler birlikte değerlendirilecektir. Morfolojik göstergeler kentsel dokunun fiziksel yapısını, erişilebilirlik göstergeleri gündelik hizmetlere ve güvenli toplanma alanlarına ulaşımı, mikroiklim göstergeleri ise gölge, açıklık ve ısı baskısı gibi çevresel koşulları temsil edecektir. PlanX Urban Resilience çıktıları, farklı doku örneklerinin dirençlilik kapasitesini ve iklim uyum önceliklerini karşılaştırmak için ortak bir gösterge dili sağlayacaktır.

Pilot analizler tamamlandığında her örnek alan için morfolojik profil, erişilebilirlik profili, mikroiklim maruziyet göstergeleri, baskın zayıflık sürücüleri ve uyum önceliği sınıfları üretilecektir. Çalışma ayrıca doku tipleri arasında hangi göstergelerin ayrıştırıcı olduğunu, hangi alanlarda fiziksel yoğunluk ile erişilebilirlik avantajlarının birlikte görüldüğünü ve hangi alanlarda sosyal kırılganlık ile çevresel maruziyetin üst üste geldiğini tartışmayı hedeflemektedir. Bu çıktıların, mahalle ölçeğindeki dirençlilik okumalarını sokak ve yapı adası ölçeğindeki kentsel tasarım kararlarıyla ilişkilendirmesi beklenmektedir. Değerlendirme, nihai bir risk hükmü üretmekten çok, farklı kentsel dokuların hangi tasarım ve planlama sorunları üzerinden uyum müdahalesi gerektirdiğini görünür kılmaya odaklanacaktır.

Çalışmanın özgün katkısı, iklim dirençliliği tartışmasını yalnızca risk katmanlarının üst üste bindirilmesi olarak değil, kentsel formun gündelik mekânsal performansıyla ilişkili bir tasarım ve planlama problemi olarak ele almasıdır. Bu yaklaşım, açık kaynaklı ve tekrarlanabilir bir CBS iş akışı içinde sokak ölçeği ile mahalle ölçeği arasında köprü kurmayı hedeflemektedir. Beklenen sonuçlar, İzmir Körfezi çevresindeki farklı kentsel dokuların hangi morfolojik ve işlevsel koşullar altında daha yüksek uyum baskısı taşıdığını ortaya koyacak; yerel yönetimler, kentsel tasarım öğrencileri ve planlama araştırmacıları için örneklenebilir bir doku okuma ve dirençlilik değerlendirme çerçevesi sunacaktır.
""".strip()


ENGLISH_ABSTRACT = """
Climate-resilient urban planning cannot be reduced to the spatial distribution of hazards. It also depends on the urban tissues through which these hazards are experienced, absorbed, or amplified. Street permeability, block size, building density, openness, shade and solar exposure, access to amenities, and social vulnerability indicators all shape how different urban fragments respond to climate stress. However, many urban resilience assessments remain at the level of administrative units, generalized risk indices, or single-hazard layers, leaving design-scale differences between urban fabrics insufficiently visible. This limitation is especially relevant in metropolitan settings such as Izmir, where coastal districts, hillside fabrics, historical centers, industrial edges, and new development areas coexist within the same urban region.

This study proposes a street-based urban morphology and resilience analytics workflow for selected urban tissue samples around the Izmir Gulf. The central research question asks how different street-fabric types vary in terms of accessibility, microclimate exposure proxies, social vulnerability, and climate adaptation priority. The planned sample set may include historical central fabric, planned grid areas, medium- and high-density apartment blocks, waterfront transformation areas, industrial and logistics edges, hillside residential fabrics, and peripheral development areas. The samples will be normalized through 400 m or 800 m walking catchments, street-corridor buffers, or equal-area morphology grids. This normalization will allow urban fragments with different sizes and spatial positions to be compared through a shared analytical frame. Sample selection will represent not only spatial distribution, but also tissue character, including street continuity, block logic, open-space relations, and everyday access potential.

The method is designed as an open-source QGIS workflow narrated through PlanX Urban Resilience. The analytical chain will combine street-network preparation, segment-based space syntax measures, network centrality, building-form metrics, morphological tessellation, Spacematrix density indicators, street-orientation structure, multi-amenity accessibility, heat-island risk, social vulnerability, and emergency accessibility. Morphological indicators will describe the physical structure of each tissue; accessibility indicators will represent access to daily amenities and emergency assembly areas; and microclimate proxies will capture environmental conditions such as shade, openness, and heat exposure. PlanX Urban Resilience outputs will provide a shared indicator language for comparing the resilience capacity and adaptation priority of different tissue samples.

Once pilot analyses are completed, the workflow will generate a morphological profile, an accessibility profile, microclimate exposure indicators, dominant weakness drivers, and adaptation-priority classes for each sample area. The study will also examine which indicators differentiate tissue types, where physical density and accessibility advantages coincide, and where social vulnerability overlaps with environmental exposure. These outputs are expected to connect neighborhood-scale resilience readings with street- and block-scale urban design decisions.

The contribution of the study is to treat climate resilience not only as an overlay of risk layers, but also as an urban design and planning question linked to the spatial performance of urban form. By connecting street-scale morphology with neighborhood-scale resilience indicators in a reproducible open-source GIS environment, the study aims to provide a transferable framework for reading climate adaptation priorities across contrasting urban fabrics. The expected outputs will help identify which morphological and functional conditions increase adaptation pressure around the Izmir Gulf and will offer a replicable assessment model for local governments, urban design students, and planning researchers.
""".strip()


def word_count(text: str) -> int:
    return len(re.findall(r"\b[\wÇĞİÖŞÜçğıöşü'-]+\b", text, flags=re.UNICODE))


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color: str = BORDER, size: str = "6") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_geometry(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = widths[idx]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")

            tc_mar = tc_pr.find(qn("w:tcMar"))
            if tc_mar is None:
                tc_mar = OxmlElement("w:tcMar")
                tc_pr.append(tc_mar)
            for side, val in (("top", 80), ("bottom", 80), ("start", 120), ("end", 120)):
                el = tc_mar.find(qn(f"w:{side}"))
                if el is None:
                    el = OxmlElement(f"w:{side}")
                    tc_mar.append(el)
                el.set(qn("w:w"), str(val))
                el.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_border(cell)


def set_run_font(run, name="Calibri", size=None, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_paragraph_spacing(paragraph, before=0, after=6, line=1.10):
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line


def add_paragraph(doc, text="", *, style=None, bold=False, italic=False, color=INK, size=11, align=None, after=6):
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    set_paragraph_spacing(p, after=after)
    if text:
        r = p.add_run(text)
        set_run_font(r, size=size, color=color, bold=bold, italic=italic)
    return p


def add_heading(doc, text: str, level: int):
    p = doc.add_paragraph(style=f"Heading {level}")
    run = p.add_run(text)
    if level == 1:
        set_run_font(run, size=16, color=BLUE, bold=True)
        set_paragraph_spacing(p, before=18, after=10)
    elif level == 2:
        set_run_font(run, size=13, color=BLUE, bold=True)
        set_paragraph_spacing(p, before=12, after=6)
    else:
        set_run_font(run, size=12, color=DARK_BLUE, bold=True)
        set_paragraph_spacing(p, before=8, after=4)
    return p


def add_bullet(doc, text: str):
    p = doc.add_paragraph(style="List Bullet")
    set_paragraph_spacing(p, after=4, line=1.208)
    run = p.add_run(text)
    set_run_font(run, size=11, color=INK)
    return p


def add_callout(doc, label: str, text: str):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, CALLOUT_FILL)
    p = cell.paragraphs[0]
    set_paragraph_spacing(p, after=2, line=1.10)
    r1 = p.add_run(f"{label}: ")
    set_run_font(r1, size=10.5, color=DARK_BLUE, bold=True)
    r2 = p.add_run(text)
    set_run_font(r2, size=10.5, color=INK)
    doc.add_paragraph()


def add_kv_table(doc, rows, widths=(2200, 7160), header=None):
    row_count = len(rows) + (1 if header else 0)
    table = doc.add_table(rows=row_count, cols=2)
    set_table_geometry(table, list(widths))
    idx = 0
    if header:
        table.cell(0, 0).merge(table.cell(0, 1))
        cell = table.cell(0, 0)
        set_cell_shading(cell, LIGHT_FILL)
        p = cell.paragraphs[0]
        set_paragraph_spacing(p, after=0)
        r = p.add_run(header)
        set_run_font(r, size=10.5, color=DARK_BLUE, bold=True)
        idx = 1
    for label, value in rows:
        c0 = table.cell(idx, 0)
        c1 = table.cell(idx, 1)
        set_cell_shading(c0, LIGHT_FILL)
        for cell in (c0, c1):
            p = cell.paragraphs[0]
            set_paragraph_spacing(p, after=0, line=1.10)
        r0 = c0.paragraphs[0].add_run(label)
        set_run_font(r0, size=10, color=DARK_BLUE, bold=True)
        r1 = c1.paragraphs[0].add_run(value)
        set_run_font(r1, size=10, color=INK)
        idx += 1
    doc.add_paragraph()
    return table


def add_text_block(doc, text: str):
    for para in text.split("\n\n"):
        p = add_paragraph(doc, para.strip(), after=7)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def configure_document(doc: Document):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for style_name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_paragraph_spacing(header, after=0)
    r = header.add_run("ICUS 2026 | Öğrenci Görüş Paketi")
    set_run_font(r, size=9, color=MUTED)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(footer, after=0)
    r = footer.add_run("Taslak belge | 13 Haziran 2026")
    set_run_font(r, size=9, color=MUTED)


def build():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_document(doc)

    # Memo masthead style opening.
    add_paragraph(doc, "Özet Bildiri Değerlendirme Paketi", bold=True, color=BLUE, size=23, after=4)
    add_paragraph(
        doc,
        "ICUS 2026 / 11. Kent Araştırmaları Kongresi için öğrenci görüş taslağı",
        color=MUTED,
        size=13,
        after=14,
    )
    add_kv_table(
        doc,
        [
            ("Çalışma yönü", "Sokak bazlı kentsel doku analitiği, kentsel morfoloji ve iklim dirençliliği"),
            ("Kongre", "11. Kent Araştırmaları Kongresi / ICUS 2026, Ankara, 12-14 Ekim 2026"),
            ("Ana tema", "İklim Değişikliğine Dirençli Kentler"),
            ("Teslim tipi", "Yalnızca özet bildiri. Tam metin bu kongreye gönderilmeyecek; genişletilmiş makale daha sonra ayrı bir dergiye hedeflenecek."),
            ("Özet son tarihi", "29 Haziran 2026"),
            ("Sunum", "Kongre sunumunu Halil Topçu yapacak şekilde planlanmaktadır."),
        ],
        header="Belge amacı",
    )

    add_callout(
        doc,
        "Kısa öneri",
        "Çalışmayı eklenti tanıtımı olarak değil, İzmir Körfezi çevresindeki farklı kentsel dokuların iklim dirençliliği açısından karşılaştırılması olarak konumlandırmak daha güçlüdür.",
    )

    add_heading(doc, "1. Yazar ve Sunum Bilgisi", 1)
    add_kv_table(
        doc,
        [
            ("1. yazar", "Yusuf Eminoğlu"),
            ("E-posta", "yusuf.eminoglu@deu.edu.tr"),
            ("ORCID", "https://orcid.org/0009-0005-6000-2934"),
            ("Kurum / görev", "Araştırma Görevlisi ve Doktora Adayı, Şehir ve Bölge Planlama Bölümü, Dokuz Eylül Üniversitesi, İzmir, Türkiye"),
            ("2. yazar / sunucu", "Halil Topçu"),
            ("E-posta", "halil.topcu2001@hotmail.com"),
            ("ORCID", "https://orcid.org/0009-0009-3366-179X"),
            ("Kurum / program", "Yüksek Lisans Öğrencisi, İzmir Demokrasi Üniversitesi, Fen Bilimleri Enstitüsü, Kentsel Tasarım Programı"),
        ],
        widths=(2500, 6860),
    )
    add_paragraph(
        doc,
        "Not: CMT sistemindeki kör değerlendirme alanlarında yazar bilgisi özet gövdesine yazılmamalıdır. Bu bilgiler yalnızca sistemdeki yazar alanlarına girilmelidir.",
        italic=True,
        color=MUTED,
        after=10,
    )

    add_heading(doc, "2. Çalışma Konumu", 1)
    add_text_block(
        doc,
        "Bu özet bildiri, iklim dirençliliğini yalnızca tehlike katmanları veya idari birim skorları üzerinden değil, kentsel doku ve sokak ağı ölçeğinde okumayı hedeflemektedir. Ana fikir, İzmir Körfezi çevresindeki farklı doku tiplerinin morfolojik, erişilebilirlik ve dirençlilik göstergeleri üzerinden karşılaştırılmasıdır.\n\n"
        "Teknik iş akışı PlanX Urban Resilience şemsiyesi altında anlatılacaktır. Ancak analiz omurgası, gerektiğinde PlanX ana eklentisinin morfoloji, space syntax, ağ merkeziliği, Spacematrix, erişilebilirlik ve mikroiklim araçlarıyla desteklenecektir. Bu ifade biçimi çalışmayı yazılım tanıtımı olmaktan çıkarır ve kentsel tasarım ile planlama araştırması olarak konumlandırır."
    )

    add_heading(doc, "3. Önerilen Örnek Alan Mantığı", 1)
    add_paragraph(doc, "Halil’in özellikle değerlendirmesi istenen kısım örnek alan seçimi ve doku tipolojisidir.", italic=True, color=MUTED)
    for item in [
        "Tarihî merkez / geleneksel yoğun doku: Kemeraltı ve çevresi gibi alanlar.",
        "Planlı ızgara ve yüksek erişilebilirlik dokusu: Alsancak, Karşıyaka veya Bostanlı çevresi.",
        "Orta-yüksek yoğunluklu apartman blok dokusu: Göztepe, Hatay, Bayraklı veya benzeri kesitler.",
        "Kıyı dönüşüm ve karma kullanım dokusu: Bayraklı-Salhane hattı gibi dönüşen parçalar.",
        "Sanayi ve lojistik kenar: Çiğli, Atatürk OSB veya liman/arka alanları.",
        "Yamaç veya kademeli gelişmiş konut dokusu: Kadifekale çevresi, Bornova yamaçları veya benzer örnekler.",
        "Periferik yeni gelişme alanı: Körfezle ilişkili yeni konut/gelişme kesitleri.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "4. CMT’ye Uygun Taslak Başlıklar", 1)
    add_heading(doc, "Türkçe başlık", 2)
    add_paragraph(
        doc,
        "Sokak Bazlı Kentsel Doku Analitiği ile İklim Dirençliliği: İzmir Körfezi Örneğinde Açık Kaynaklı Bir QGIS İş Akışı",
        bold=True,
        color=INK,
    )
    add_heading(doc, "English title", 2)
    add_paragraph(
        doc,
        "Street-Based Urban Tissue Analytics for Climate Resilience: An Open-Source QGIS Workflow in the Izmir Gulf",
        bold=True,
        color=INK,
    )

    add_heading(doc, "5. Türkçe Özet Taslağı", 1)
    add_paragraph(doc, f"Kelime sayısı: {word_count(TURKISH_ABSTRACT)}", italic=True, color=MUTED, after=4)
    add_text_block(doc, TURKISH_ABSTRACT)

    add_heading(doc, "Türkçe Anahtar Kelimeler", 2)
    add_paragraph(doc, "kentsel dirençlilik; kentsel morfoloji; kentsel doku; sokak ağı; açık kaynak CBS")

    doc.add_section(WD_SECTION.NEW_PAGE)
    configure_document(doc)

    add_heading(doc, "6. English Abstract Draft", 1)
    add_paragraph(doc, f"Word count: {word_count(ENGLISH_ABSTRACT)}", italic=True, color=MUTED, after=4)
    add_text_block(doc, ENGLISH_ABSTRACT)

    add_heading(doc, "English Keywords", 2)
    add_paragraph(doc, "urban resilience; urban morphology; urban fabric; street networks; open-source GIS")

    add_heading(doc, "7. Halil Topçu’dan Beklenen Görüş", 1)
    for item in [
        "Örnek doku aileleri İzmir Körfezi için doğru ve savunulabilir mi?",
        "Kentsel tasarım açısından hangi 5-7 örnek alan daha temsilî olur?",
        "Analiz birimi 400/800 m yürüme yakalaması mı, 500 m grid mi, yoksa sokak koridoru mu olmalı?",
        "Başlık çalışma fikrini doğru taşıyor mu?",
        "Özet metni öğrencinin sunabileceği kadar anlaşılır mı, yoksa yöntem kısmı sadeleştirilmeli mi?",
        "Kongrede sunum yapılırken hangi harita/şema dili daha iyi çalışır?",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "8. Kongre Uyum Kontrolü", 1)
    add_kv_table(
        doc,
        [
            ("Özet uzunluğu", f"Türkçe özet {word_count(TURKISH_ABSTRACT)} kelime. CMT alanına göre gerekirse son sıkıştırma yapılacak."),
            ("Dil koşulu", "Türkçe ve İngilizce başlık, özet ve anahtar kelimeler hazırlanmıştır."),
            ("Yazar bilgisi", "Özet gövdesinde yazar bilgisi yoktur; bilgiler yalnızca metadata kısmındadır."),
            ("Kaynakça / tablo / şekil", "Özet gövdesinde kaynakça, tablo veya şekil yoktur."),
            ("Kongre teması", "İklim dirençliliği, kentsel kırılganlık, mekânsal adalet, yerel planlama ve kentsel ekosistemler temalarıyla uyumludur."),
            ("Tam metin kararı", "Bu kongreye tam metin gönderilmeyecek; genişletilmiş makale başka bir dergiye hazırlanacaktır."),
        ],
        widths=(2300, 7060),
    )

    doc.save(OUT_DOCX)


if __name__ == "__main__":
    build()
    print(OUT_DOCX)
